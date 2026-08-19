from io import BytesIO

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def create_tailored_resume_docx(result: dict) -> BytesIO:
    """
    Generate a professional editable DOCX resume from the
    structured JSON returned by the Resume Tailor AI.
    """

    document = Document()

    # ---------------------------------------------------------
    # PAGE SETUP
    # ---------------------------------------------------------

    section = document.sections[0]

    section.top_margin = Pt(36)
    section.bottom_margin = Pt(36)
    section.left_margin = Pt(50)
    section.right_margin = Pt(50)

    # ---------------------------------------------------------
    # DEFAULT FONT
    # ---------------------------------------------------------

    styles = document.styles

    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(9.5)

    # ---------------------------------------------------------
    # NAME
    # ---------------------------------------------------------

    name = result.get("name", "Candidate")

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = paragraph.add_run(name)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(18)

    # ---------------------------------------------------------
    # CONTACT INFORMATION
    # ---------------------------------------------------------

    contact = result.get("contact", {})

    contact_parts = []

    for key in [
        "location",
        "email",
        "phone",
        "github",
        "linkedin"
    ]:
        value = contact.get(key)

        if value:
            value = str(value).strip()

            if value:
                contact_parts.append(value)

    if contact_parts:

        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = paragraph.add_run(
            " | ".join(contact_parts)
        )

        run.font.name = "Arial"
        run.font.size = Pt(9)

    # ---------------------------------------------------------
    # PROFESSIONAL SUMMARY
    # ---------------------------------------------------------

    summary = result.get(
        "professional_summary",
        ""
    )

    if summary:

        add_heading(
            document,
            "PROFESSIONAL SUMMARY"
        )

        paragraph = document.add_paragraph(
            str(summary).strip()
        )

        paragraph.paragraph_format.space_after = Pt(5)

    # ---------------------------------------------------------
    # TECHNICAL SKILLS
    # ---------------------------------------------------------

    technical_skills = result.get(
        "technical_skills",
        {}
    )

    if technical_skills:

        add_heading(
            document,
            "TECHNICAL SKILLS"
        )

        for category, skills in technical_skills.items():

            if not skills:
                continue

            if isinstance(skills, str):
                skills = [skills]

            skills = [
                str(skill).strip()
                for skill in skills
                if str(skill).strip()
            ]

            if not skills:
                continue

            paragraph = document.add_paragraph()

            paragraph.paragraph_format.space_after = Pt(2)

            label = paragraph.add_run(
                f"{category}: "
            )

            label.bold = True

            paragraph.add_run(
                ", ".join(skills)
            )

    # ---------------------------------------------------------
    # PROJECTS
    # ---------------------------------------------------------

    projects = result.get(
        "projects",
        []
    )

    if projects:

        add_heading(
            document,
            "PROJECTS"
        )

        for project in projects:

            if not isinstance(project, dict):
                continue

            title = project.get(
                "title",
                ""
            )

            technologies = project.get(
                "technologies",
                ""
            )

            # Project title
            if title:

                paragraph = document.add_paragraph()

                paragraph.paragraph_format.space_before = Pt(4)
                paragraph.paragraph_format.space_after = Pt(1)

                run = paragraph.add_run(
                    str(title).strip()
                )

                run.bold = True
                run.font.size = Pt(10.5)

            # Technologies
            if technologies:

                paragraph = document.add_paragraph()

                paragraph.paragraph_format.space_after = Pt(2)

                run = paragraph.add_run(
                    str(technologies).strip()
                )

                run.italic = True
                run.font.size = Pt(9)

            # Project bullets
            bullets = project.get(
                "bullets",
                []
            )

            for bullet in bullets:

                if not bullet:
                    continue

                paragraph = document.add_paragraph(
                    style="List Bullet"
                )

                paragraph.paragraph_format.left_indent = Pt(12)
                paragraph.paragraph_format.space_after = Pt(2)

                paragraph.add_run(
                    str(bullet).strip()
                )

    # ---------------------------------------------------------
    # EDUCATION
    # ---------------------------------------------------------

    education = result.get(
        "education",
        []
    )

    if education:

        add_heading(
            document,
            "EDUCATION"
        )

        for item in education:

            if not isinstance(item, dict):
                continue

            degree = item.get(
                "degree",
                ""
            )

            institution = item.get(
                "institution",
                ""
            )

            details = item.get(
                "details",
                ""
            )

            if degree:

                paragraph = document.add_paragraph()

                paragraph.paragraph_format.space_after = Pt(1)

                run = paragraph.add_run(
                    str(degree).strip()
                )

                run.bold = True
                run.font.size = Pt(10.5)

            if institution:

                paragraph = document.add_paragraph()

                paragraph.paragraph_format.space_after = Pt(1)

                paragraph.add_run(
                    str(institution).strip()
                )

            if details:

                paragraph = document.add_paragraph()

                paragraph.paragraph_format.space_after = Pt(3)

                paragraph.add_run(
                    str(details).strip()
                )

    # ---------------------------------------------------------
    # LEADERSHIP & COMMUNITY EXPERIENCE
    # ---------------------------------------------------------

    leadership = result.get(
        "leadership",
        []
    )

    if leadership:

        add_heading(
            document,
            "LEADERSHIP & COMMUNITY EXPERIENCE"
        )

        for role in leadership:

            if not isinstance(role, dict):
                continue

            title = role.get(
                "title",
                ""
            )

            organization = role.get(
                "organization",
                ""
            )

            if title:

                paragraph = document.add_paragraph()

                paragraph.paragraph_format.space_after = Pt(1)

                run = paragraph.add_run(
                    str(title).strip()
                )

                run.bold = True
                run.font.size = Pt(10.5)

            if organization:

                paragraph = document.add_paragraph()

                paragraph.paragraph_format.space_after = Pt(2)

                paragraph.add_run(
                    str(organization).strip()
                )

            bullets = role.get(
                "bullets",
                []
            )

            for bullet in bullets:

                if not bullet:
                    continue

                paragraph = document.add_paragraph(
                    style="List Bullet"
                )

                paragraph.paragraph_format.left_indent = Pt(12)
                paragraph.paragraph_format.space_after = Pt(2)

                paragraph.add_run(
                    str(bullet).strip()
                )

    # ---------------------------------------------------------
    # CERTIFICATIONS
    # ---------------------------------------------------------

    certifications = result.get(
        "certifications",
        []
    )

    if certifications:

        add_heading(
            document,
            "CERTIFICATIONS"
        )

        for certification in certifications:

            if not certification:
                continue

            paragraph = document.add_paragraph(
                style="List Bullet"
            )

            paragraph.paragraph_format.left_indent = Pt(12)
            paragraph.paragraph_format.space_after = Pt(2)

            paragraph.add_run(
                str(certification).strip()
            )

    # ---------------------------------------------------------
    # ACHIEVEMENTS
    # ---------------------------------------------------------

    achievements = result.get(
        "achievements",
        []
    )

    if achievements:

        add_heading(
            document,
            "ACHIEVEMENTS"
        )

        for achievement in achievements:

            if not achievement:
                continue

            paragraph = document.add_paragraph(
                style="List Bullet"
            )

            paragraph.paragraph_format.left_indent = Pt(12)
            paragraph.paragraph_format.space_after = Pt(2)

            paragraph.add_run(
                str(achievement).strip()
            )

    # ---------------------------------------------------------
    # SAVE DOCX TO MEMORY
    # ---------------------------------------------------------

    file_stream = BytesIO()

    document.save(file_stream)

    file_stream.seek(0)

    return file_stream


def add_heading(document, text):
    """
    Add a clean resume section heading.
    """

    paragraph = document.add_paragraph()

    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(3)

    run = paragraph.add_run(text)

    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(11)

    return paragraph